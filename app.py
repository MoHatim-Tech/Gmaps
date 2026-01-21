import streamlit as st
import pandas as pd
from playwright.sync_api import sync_playwright
import time
import re
import requests
from bs4 import BeautifulSoup
import asyncio
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import subprocess

# وظيفة لتثبيت متصفح Playwright تلقائياً على السيرفر
def install_playwright_browsers():
    try:
        subprocess.run(["playwright", "install", "chromium"], check=True)
    except Exception as e:
        st.error(f"Error installing browsers: {e}")

# تنفيذ التثبيت عند بدء التطبيق (للمنصات السحابية مثل Streamlit Cloud)
if 'playwright_installed' not in st.session_state:
    with st.spinner("جاري تهيئة محرك البحث (قد يستغرق دقيقة في المرة الأولى)..."):
        install_playwright_browsers()
        st.session_state['playwright_installed'] = True

# حل مشكلة NotImplementedError على ويندوز
if sys.platform == 'win32':
    asyncio.set_event_loop_policy(asyncio.WindowsProactorEventLoopPolicy())

def extract_emails_from_url(url):
    try:
        if not url or url == "N/A": return "N/A"
        if not url.startswith('http'): url = 'https://' + url
        
        # محاولة البحث في عدة صفحات شائعة
        targets = [url]
        base_url = "/".join(url.split("/")[:3])
        # قائمة الصفحات المحتملة للاتصال
        potential_pages = ['contact', 'contact-us', 'about', 'about-us', 'support', 'terms']
        for page_name in potential_pages:
            targets.append(f"{base_url}/{page_name}")
            targets.append(f"{base_url}/ar/{page_name}") # دعم المواقع العربية
            targets.append(f"{base_url}/en/{page_name}") # دعم المواقع الإنجليزية
        
        # إزالة التكرار مع الحفاظ على الترتيب
        targets = list(dict.fromkeys(targets))
        
        all_emails = set()
        for target in targets:
            try:
                response = requests.get(target, timeout=5, headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'})
                if response.status_code == 200:
                    # استخراج الإيميلات مع استبعاد الصور والملفات الشائعة
                    found = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', response.text)
                    for email in found:
                        if not any(ext in email.lower() for ext in ['.png', '.jpg', '.jpeg', '.gif', '.svg', 'wix', 'sentry']):
                            all_emails.add(email)
            except: continue
            if all_emails: break
            
        return ", ".join(list(all_emails)) if all_emails else "N/A"
    except:
        return "N/A"

def scrape_google_maps(search_query, max_results=10, data_placeholder=None, progress_bar=None):
    with sync_playwright() as p:
        results = []
        try:
            # استخدام إعدادات متقدمة لتجنب الاكتشاف ومحاكاة إنسان
            browser = p.chromium.launch(headless=True, args=[
                "--no-sandbox",
                "--disable-setuid-sandbox",
                "--disable-blink-features=AutomationControlled"
            ])
            context = browser.new_context(
                user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36",
                locale="ar-SA",
                viewport={'width': 1920, 'height': 1080}
            )
            page = context.new_page()
            
            # منع اكتشاف المتصفح كآلي
            page.add_init_script("Object.defineProperty(navigator, 'webdriver', {get: () => undefined})")
            
            # 1. الذهاب للموقع الرئيسي أولاً لتضليل أنظمة الكشف
            st.toast("🔍 جاري الدخول إلى خرائط جوجل...")
            page.goto("https://www.google.com/maps?hl=ar", wait_until="domcontentloaded", timeout=60000)
            time.sleep(4)
            
            # محاولة تخطي أي نوافذ منبثقة للموافقة
            try:
                consent_btn = page.locator('button:has-text("قبول"), button:has-text("وافق"), button:has-text("Accept")').first
                if consent_btn.is_visible():
                    consent_btn.click()
                    time.sleep(2)
            except: pass

            # 2. استخدام مربع البحث بدلاً من الرابط المباشر
            st.toast(f"📝 جاري البحث عن: {search_query}")
            search_box = page.locator('#searchboxinput')
            search_box.wait_for(state="visible", timeout=20000)
            search_box.fill(search_query)
            page.keyboard.press("Enter")
            
            # الانتظار حتى تحميل النتائج الأولية
            time.sleep(6)

            seen_names = set()
            scroll_attempts = 0
            max_scroll_attempts = 60 
            
            while len(results) < max_results and scroll_attempts < max_scroll_attempts:
                # التحقق مما إذا كانت هناك نتيجة واحدة مباشرة (صفحة مؤسسة مفتوحة)
                if page.locator('h1.DUwDvf').count() > 0:
                    name = page.locator('h1.DUwDvf').first.inner_text()
                    if name not in seen_names:
                        address = page.locator('button[data-item-id="address"]').first.inner_text() if page.locator('button[data-item-id="address"]').count() > 0 else "N/A"
                        phone = page.locator('button[data-item-id^="phone:tel:"]').first.inner_text() if page.locator('button[data-item-id^="phone:tel:"]').count() > 0 else "N/A"
                        website = page.locator('a[data-item-id="authority"]').first.get_attribute('href') if page.locator('a[data-item-id="authority"]').count() > 0 else "N/A"
                        
                        results.append({
                            "🏢 اسم المؤسسة": name, "📞 رقم الهاتف": phone, "🌐 الموقع الالكتروني": website,
                            "📧 الايميل": extract_emails_from_url(website) if website != "N/A" else "N/A",
                            "📍 موقع المكتب": address
                        })
                        seen_names.add(name)
                        if data_placeholder: data_placeholder.dataframe(pd.DataFrame(results), use_container_width=True)
                        if max_results == 1: break # إذا طلب نتيجة واحدة وجدناها

                # البحث عن عناصر النتائج في القائمة
                item_selectors = ['.Nv262d', '.hfpxzc', 'a[href*="/maps/place/"]']
                items = []
                for sel in item_selectors:
                    found = page.locator(sel).all()
                    if len(found) > 0:
                        items = found
                        break
                
                if not items:
                    page.mouse.wheel(0, 2000)
                    time.sleep(3)
                    scroll_attempts += 1
                    continue

                for item in items:
                    if len(results) >= max_results: break
                    try:
                        # الحصول على الاسم الأولي للتحقق
                        name_text = item.get_attribute("aria-label") or item.inner_text().split('\n')[0]
                        if not name_text or name_text in seen_names: continue

                        item.scroll_into_view_if_needed()
                        item.click(force=True, timeout=10000)
                        time.sleep(2) # انتظار تحميل التفاصيل
                        
                        # استخراج البيانات من اللوحة الجانبية
                        name_loc = page.locator('h1.DUwDvf')
                        if name_loc.count() > 0:
                            side_name = name_loc.first.inner_text()
                            if side_name in seen_names: continue
                            
                            address = page.locator('button[data-item-id="address"]').first.inner_text() if page.locator('button[data-item-id="address"]').count() > 0 else "N/A"
                            phone = page.locator('button[data-item-id^="phone:tel:"]').first.inner_text() if page.locator('button[data-item-id^="phone:tel:"]').count() > 0 else "N/A"
                            website = page.locator('a[data-item-id="authority"]').first.get_attribute('href') if page.locator('a[data-item-id="authority"]').count() > 0 else "N/A"
                            
                            results.append({
                                "🏢 اسم المؤسسة": side_name, "📞 رقم الهاتف": phone, "🌐 الموقع الالكتروني": website,
                                "📧 الايميل": extract_emails_from_url(website) if website != "N/A" else "N/A",
                                "📍 موقع المكتب": address
                            })
                            seen_names.add(side_name)
                            if progress_bar: progress_bar.progress(len(results) / max_results)
                            if data_placeholder: data_placeholder.dataframe(pd.DataFrame(results), use_container_width=True)
                    except: continue
                
                # التمرير لأسفل لتحميل المزيد
                try:
                    feed = page.locator('div[role="feed"]')
                    if feed.count() > 0:
                        feed.evaluate("el => el.scrollBy(0, 4000)")
                    else:
                        page.mouse.wheel(0, 4000)
                except: page.mouse.wheel(0, 4000)
                
                time.sleep(3)
                scroll_attempts += 1
                if "reached the end" in page.content() or "نهاية القائمة" in page.content(): break
                    
            browser.close()
            return results
        except Exception as e:
            if 'browser' in locals(): browser.close()
            return results

# إعدادات الواجهة
st.set_page_config(page_title="مستخرج بيانات خرائط جوجل", layout="wide", initial_sidebar_state="expanded")

# تصميم عصري وأنيق مع تجاوز تنسيقات Streamlit
style_code = """<link href="https://fonts.googleapis.com/css2?family=Tajawal:wght@400;500;700&display=swap" rel="stylesheet"><style>
    body, .stApp { font-family: 'Tajawal', sans-serif !important; direction: RTL !important; text-align: right !important; background-color: #F8FAFC !important; }
    h1, h2, h3, p, span, label { font-family: 'Tajawal', sans-serif !important; text-align: right !important; color: #1E3A8A !important; }
    
    /* تنسيق شريط التمرير الجانبي */
    [data-testid="stSidebar"] {
        background-color: #FFFFFF !important;
        border-left: 1px solid #E2E8F0 !important;
    }
    
    [data-testid="stSidebar"] .stMarkdown h3 {
        color: #2563EB !important;
        border-bottom: 2px solid #F1F5F9;
        padding-bottom: 10px;
    }

    /* تنسيق مربعات الإدخال */
    .stTextInput div[data-baseweb="input"], .stNumberInput div[data-baseweb="input"] {
        border: 1px solid #CBD5E1 !important;
        border-radius: 8px !important;
        background-color: white !important;
    }
    
    /* تنسيق جدول النتائج - LTR */
    [data-testid="stDataFrame"], [data-testid="stTable"] {
        direction: LTR !important;
        text-align: left !important;
        background-color: white !important;
        border-radius: 12px !important;
        box-shadow: 0 1px 3px 0 rgba(0, 0, 0, 0.1) !important;
    }
    
    /* إخفاء الرسائل المزعجة */
    [data-testid="stInputHelperText"], .st-emotion-cache-1pxm8v5, .st-emotion-cache-10trblm { display: none !important; }
    
    /* تنسيق الزر */
    .stButton button {
        background-color: #2563EB !important;
        color: white !important;
        border-radius: 8px !important;
        padding: 0.6rem 1rem !important;
        width: 100% !important;
        font-weight: bold !important;
        font-family: 'Tajawal', sans-serif !important;
        border: none !important;
        transition: all 0.2s ease !important;
    }

    .stButton button p { color: white !important; }
    
    .stButton button:hover {
        background-color: #1E40AF !important;
        box-shadow: 0 4px 12px rgba(37, 99, 235, 0.2) !important;
    }
    
    # MainMenu, footer, header { visibility: hidden !important; }
    
    .developer-footer {
        position: fixed;
        bottom: 0;
        left: 0;
        width: 100%;
        background-color: #1E3A8A;
        color: white;
        text-align: center;
        padding: 8px 0;
        font-family: 'Tajawal', sans-serif;
        z-index: 100;
        font-size: 0.9rem;
    }
    </style>"""
st.markdown(style_code, unsafe_allow_html=True)

# إضافة حقوق المطور في الأسفل
st.markdown("""
    <div class="developer-footer">
        👨‍💻 تطوير: <b>عبدالمنعم حاتم</b> | 📞: +966544451878 | 📧: info@mohatim.tech
    </div>
    """, unsafe_allow_html=True)

st.title("🔍 نظام استخراج البيانات الذكي")

# تنظيم المدخلات في الشريط الجانبي
with st.sidebar:
    st.markdown("### 🛠️ إعدادات البحث")
    business_type = st.text_input("مجال المؤسسة", placeholder="مطاعم، فنادق...")
    city = st.text_input("المدينة", placeholder="الرياض، دبي...")
    country = st.text_input("الدولة", placeholder="السعودية...")
    max_res = st.number_input("عدد النتائج المطلوبة", min_value=1, max_value=500, value=10, step=1)
    
    st.markdown("---")
    search_clicked = st.button("🚀 ابدأ عملية الاستخراج")
    
    st.markdown("### 📖 تعليمات الاستخدام")
    st.info("""
    1. أدخل نوع النشاط التجاري.
    2. حدد المدينة والدولة بدقة.
    3. اختر عدد النتائج (الحد الأقصى 500).
    4. اضغط على زر البدء وانتظر النتائج.
    5. يمكنك تحميل البيانات بصيغة Word أو CSV.
    """)

def create_word_doc(data):
    doc = Document()
    
    # إعداد النمط الافتراضي لدعم العربية
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    doc.add_heading('نتائج استخراج بيانات خرائط جوجل', 0).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    for entry in data:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"🏢 اسم المؤسسة: ").bold = True
        p.add_run(f"{entry['🏢 اسم المؤسسة']}\n")
        
        p.add_run(f"📞 رقم الهاتف: ").bold = True
        p.add_run(f"{entry['📞 رقم الهاتف']}\n")
        
        p.add_run(f"🌐 الموقع الالكتروني: ").bold = True
        p.add_run(f"{entry['🌐 الموقع الالكتروني']}\n")
        
        p.add_run(f"📧 الايميل: ").bold = True
        p.add_run(f"{entry['📧 الايميل']}\n")
        
        p.add_run(f"📍 موقع المكتب: ").bold = True
        p.add_run(f"{entry['📍 موقع المكتب']}\n")
        
        doc.add_paragraph("-" * 30).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # إضافة حقوق المطور في نهاية الملف
    doc.add_paragraph("\n")
    dev_info = doc.add_paragraph()
    dev_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dev_info.add_run("تم الاستخراج بواسطة نظام استخراج البيانات الذكي")
    run.font.size = Pt(10)
    run.italic = True
    
    dev_contact = doc.add_paragraph()
    dev_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = dev_contact.add_run("تطوير: عبدالمنعم حاتم | جوال: +966544451878 | ايميل: info@mohatim.tech")
    run2.font.size = Pt(10)
    run2.bold = True

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

st.markdown("<br>", unsafe_allow_html=True)
if search_clicked:
    # التحقق من وجود قيمة واحدة على الأقل للبحث
    if business_type or city or country:
        # بناء نص البحث بشكل أكثر دقة لضمان تصفية المدينة والدولة
        query_parts = []
        if business_type:
            query_parts.append(business_type)
        
        location_parts = []
        if city:
            location_parts.append(city)
        if country:
            location_parts.append(country)
            
        if location_parts:
            query_parts.append("في")
            query_parts.append(", ".join(location_parts))
        
        query = " ".join(query_parts)
        st.info(f"جاري البحث عن: {query}...")
        
        progress_bar = st.progress(0)
        data_placeholder = st.empty()
        
        final_data = scrape_google_maps(query, max_res, data_placeholder, progress_bar)
        
        if final_data:
            st.success("اكتملت عملية الاستخراج!")
            
            df = pd.DataFrame(final_data)
            
            # عرض أزرار التحميل في أعمدة
            col_word, col_csv = st.columns(2)
            
            with col_word:
                # إنشاء ملف Word
                word_data = create_word_doc(final_data)
                st.download_button(
                    label="تحميل النتائج كـ Word (.docx)",
                    data=word_data,
                    file_name="google_maps_results.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            
            with col_csv:
                # إنشاء ملف CSV
                csv = df.to_csv(index=False).encode('utf-8-sig')
                st.download_button(
                    label="تحميل النتائج كـ CSV (.csv)",
                    data=csv,
                    file_name="google_maps_results.csv",
                    mime="text/csv",
                    use_container_width=True
                )
        else:
            st.error("لم يتم العثور على نتائج.")
    else:
        st.warning("يرجى إدخال معلومة واحدة على الأقل للبحث (المجال أو المدينة أو الدولة).")
